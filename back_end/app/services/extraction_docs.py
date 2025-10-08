"""
Extraction Docs - Extraction de texte depuis des documents

Ce module permet d'extraire du texte depuis :
- Fichiers PDF
- Documents Word (.docx)
- Fichiers texte
"""

from typing import Optional, Dict
import PyPDF2
from docx import Document
import io

class DocumentExtractor:
    """
    Extracteur de texte depuis différents formats de documents
    """
    
    def extract_from_pdf(self, pdf_content: bytes) -> Dict:
        """
        Extrait le texte d'un PDF
        
        Args:
            pdf_content: Contenu binaire du PDF
            
        Returns:
            Dictionnaire avec le texte et métadonnées
        """
        try:
            pdf_file = io.BytesIO(pdf_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            return {
                'success': True,
                'text': text.strip(),
                'num_pages': len(pdf_reader.pages),
                'metadata': pdf_reader.metadata,
                'error': None
            }
        
        except Exception as e:
            return {
                'success': False,
                'text': None,
                'error': f'Erreur extraction PDF: {str(e)}'
            }
    
    def extract_from_docx(self, docx_content: bytes) -> Dict:
        """
        Extrait le texte d'un document Word
        
        Args:
            docx_content: Contenu binaire du fichier .docx
            
        Returns:
            Dictionnaire avec le texte
        """
        try:
            docx_file = io.BytesIO(docx_content)
            doc = Document(docx_file)
            
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            return {
                'success': True,
                'text': text.strip(),
                'num_paragraphs': len(doc.paragraphs),
                'error': None
            }
        
        except Exception as e:
            return {
                'success': False,
                'text': None,
                'error': f'Erreur extraction DOCX: {str(e)}'
            }
    
    def extract_from_file(self, file_path: str) -> Dict:
        """
        Extrait le texte d'un fichier (détecte automatiquement le type)
        
        Args:
            file_path: Chemin du fichier
            
        Returns:
            Dictionnaire avec le texte extrait
        """
        try:
            with open(file_path, 'rb') as f:
                content = f.read()
            
            if file_path.endswith('.pdf'):
                return self.extract_from_pdf(content)
            elif file_path.endswith('.docx'):
                return self.extract_from_docx(content)
            elif file_path.endswith('.txt'):
                return {
                    'success': True,
                    'text': content.decode('utf-8'),
                    'error': None
                }
            else:
                return {
                    'success': False,
                    'text': None,
                    'error': 'Format de fichier non supporté'
                }
        
        except Exception as e:
            return {
                'success': False,
                'text': None,
                'error': f'Erreur lecture fichier: {str(e)}'
            }


# Instance globale de l'extracteur
extractor = DocumentExtractor()


def extract_text_from_pdf(pdf_content: bytes) -> str:
    """
    Fonction utilitaire pour extraire du texte d'un PDF
    
    Args:
        pdf_content: Contenu binaire du PDF
        
    Returns:
        Texte extrait
    """
    result = extractor.extract_from_pdf(pdf_content)
    return result['text'] if result['success'] else None


# Test si exécuté directement
if __name__ == "__main__":
    print("=" * 70)
    print("TEST EXTRACTION DOCS")
    print("=" * 70)
    print("\nModule prêt pour extraction de PDF et DOCX")
    print("Utilisation : extractor.extract_from_pdf(pdf_bytes)")
    print("            : extractor.extract_from_docx(docx_bytes)")